"""
Génère le template DAT par défaut compatible docxtpl.
- Champs texte simple  : {{ variable }}
- Champs rich text     : {{r variable }}   (gras, italique, couleur, etc.)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_rich_para(doc, field_name):
    """Ajoute un paragraphe avec un champ rich text {{r field_name }}."""
    p = doc.add_paragraph()
    p.add_run(f'{{{{r {field_name} }}}}')
    return p


def create_dat_template():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ===== PAGE DE GARDE =====
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('DOSSIER D\'ARCHITECTURE TECHNIQUE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('{{ titre_projet }}')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 0, 145)

    doc.add_paragraph()
    doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate([
        ('Chef de Projet', '{{ chef_projet }}'),
        ('Contact Technique', '{{ contact_tech }}'),
        ('Date', '{{ date }}'),
        ('Version', '1.0'),
    ]):
        cell_l = info_table.rows[i].cells[0]
        cell_v = info_table.rows[i].cells[1]
        cell_l.text = label
        cell_v.text = value
        set_cell_shading(cell_l, 'E3E3FD')
        cell_l.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ===== 1. OBJET DU DOCUMENT =====
    doc.add_heading('1. Objet du Document', 1)
    add_rich_para(doc, 'objet_document')
    doc.add_paragraph()
    doc.add_paragraph('{{ description_doc }}')
    doc.add_paragraph()

    # ===== 2. DOCUMENTS DE RÉFÉRENCE =====
    doc.add_heading('2. Documents de Référence', 1)

    ref_table = doc.add_table(rows=1, cols=3)
    ref_table.style = 'Table Grid'
    for i, h in enumerate(['Émetteur', 'Document', 'Version']):
        cell = ref_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '000091')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('{%tr for ref in documents_reference %}')
    ref_row_para = doc.add_paragraph('{{ ref.emetteur }} | {{ ref.document }} | {{ ref.version }}')
    doc.add_paragraph('{%tr endfor %}')
    doc.add_paragraph()

    # ===== 3. ACTEURS =====
    doc.add_heading('3. Acteurs du Projet', 1)

    actor_table = doc.add_table(rows=1, cols=4)
    actor_table.style = 'Table Grid'
    for i, h in enumerate(['Acteur', 'Rôle', 'Droits', 'Commentaires']):
        cell = actor_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '000091')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('{% for acteur in acteurs %}')
    doc.add_paragraph('• {{ acteur.acteur }} — {{ acteur.role }} ({{ acteur.droits }})')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # ===== 4. SPÉCIFICATIONS FONCTIONNELLES =====
    doc.add_heading('4. Spécifications Fonctionnelles', 1)

    doc.add_heading('4.1 Description Fonctionnelle', 2)
    add_rich_para(doc, 'schema_description')
    doc.add_paragraph()

    doc.add_heading('4.2 Briques Fonctionnelles', 2)
    doc.add_paragraph('{% for b in briques_fonctionnelles %}')
    doc.add_paragraph('• {{ b.brique }} : {{ b.description }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    doc.add_heading('4.3 Échanges de Données', 2)
    echange_table = doc.add_table(rows=1, cols=5)
    echange_table.style = 'Table Grid'
    for i, h in enumerate(['Type', 'Source', 'Destination', 'Fréquence', 'HNO']):
        cell = echange_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '000091')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('{% for e in echanges_donnees %}')
    doc.add_paragraph('• {{ e.type_echange }} : {{ e.source }} → {{ e.destination }} ({{ e.frequence }})')
    doc.add_paragraph('{% endfor %}')

    doc.add_page_break()

    # ===== 5. ARCHITECTURE TECHNIQUE =====
    doc.add_heading('5. Architecture Technique', 1)

    doc.add_heading('5.1 Description Générale', 2)
    add_rich_para(doc, 'description_architecture')

    doc.add_heading('5.2 Authentification', 2)
    add_rich_para(doc, 'description_authentification')

    doc.add_heading('5.3 Administration Technique', 2)
    add_rich_para(doc, 'description_administrationtechnique')

    doc.add_heading('5.4 Administration Fonctionnelle', 2)
    add_rich_para(doc, 'description_adminfonctionnelle')

    doc.add_heading('5.5 Communication Inter-Applicative', 2)
    add_rich_para(doc, 'description_interapplicative')

    doc.add_heading('5.6 Choix Technologiques', 2)
    doc.add_paragraph('{% for tech in choix_technologiques %}')
    doc.add_paragraph('• {{ tech.tiers }} : {{ tech.produit }} v{{ tech.version }}')
    doc.add_paragraph('{% endfor %}')

    p = doc.add_paragraph()
    p.add_run('Données sensibles / DR : ').bold = True
    p.add_run('{{ segmentation_dr }}')

    doc.add_heading('5.7 Noms DNS', 2)
    doc.add_paragraph('{% for dns in dns_nom %}')
    doc.add_paragraph('• {{ dns.nom_dns }} → {{ dns.machine_associe }}')
    doc.add_paragraph('{% endfor %}')

    doc.add_page_break()

    # ===== 6. INFRASTRUCTURE =====
    doc.add_heading('6. Infrastructure', 1)

    doc.add_heading('6.1 Machines Virtuelles', 2)
    vm_table = doc.add_table(rows=1, cols=7)
    vm_table.style = 'Table Grid'
    for i, h in enumerate(['Env.', 'Nom', 'Rôle', 'OS', 'CPU', 'RAM (Go)', 'Résilience']):
        cell = vm_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '000091')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('{% for vm in vms %}')
    doc.add_paragraph('• {{ vm.nom }} ({{ vm.environnement }}) — {{ vm.role }} | {{ vm.os }} | {{ vm.cpu }} CPU | {{ vm.ram }} Go')
    doc.add_paragraph('{% endfor %}')

    # ===== 7. CYCLE DE VIE =====
    doc.add_heading('7. Cycle de Vie', 1)

    doc.add_heading('7.1 Déploiement', 2)
    add_rich_para(doc, 'deploiement')

    doc.add_heading('7.2 Migration / Reprise de Données', 2)
    add_rich_para(doc, 'migration_reprise')

    doc.add_heading('7.3 Supervision et Observabilité', 2)
    add_rich_para(doc, 'supervision')

    doc.add_heading('7.4 Sauvegardes et Restauration', 2)
    add_rich_para(doc, 'sauvegarde_restauration')

    # ===== 8. DÉPENDANCES =====
    doc.add_heading('8. Dépendances', 1)

    doc.add_heading('8.1 Dépendances sur Mon Application', 2)
    doc.add_paragraph('{% for dep in dependances_externes %}')
    doc.add_paragraph('• {{ dep.dependance }} : {{ dep.impact }}')
    doc.add_paragraph('{% endfor %}')

    doc.add_heading('8.2 Mon Application vers Services Externes', 2)
    doc.add_paragraph('{% for dep in dependance_app_externes %}')
    doc.add_paragraph('• {{ dep.name_application }} : {{ dep.description_impact }}')
    doc.add_paragraph('{% endfor %}')

    # ===== 9. CONTRAINTES ET SLA =====
    doc.add_heading('9. Contraintes et Niveau de Service', 1)

    doc.add_heading('9.1 Contraintes', 2)
    add_rich_para(doc, 'contraintes')

    doc.add_heading('9.2 Niveau de Services (SLA)', 2)
    add_rich_para(doc, 'niveau_services')

    output_path = 'app/templates/dat_template.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template créé : {output_path}")
    return output_path


if __name__ == "__main__":
    create_dat_template()
